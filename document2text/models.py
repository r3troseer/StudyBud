import docx
from django.db import models


# Create your models here.
class Document(models.Model):
    text = models.TextField()
    document = models.FileField(upload_to="documents/")
    uploaded_at = models.DateTimeField(auto_now_add=True)

    def document_to_text(self):
        print("here")
        doc = docx.Document(self.document)
        docText = "\n\n".join(paragraph.text for paragraph in doc.paragraphs)
        # print(docText)
        self.text = str(docText)
        self.save()
        return

    def __str__(self):
        return self.document.name
